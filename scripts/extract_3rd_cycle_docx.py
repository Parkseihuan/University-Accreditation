from docx import Document
import os

DOCX_PATH = "3주기 - 대학자체진단평가보고서_ 교원 및 직원.docx"
OUTPUT_TXT = "3rd_cycle_extracted.txt"

def extract_content():
    if not os.path.exists(DOCX_PATH):
        print(f"File not found: {DOCX_PATH}")
        return

    doc = Document(DOCX_PATH)
    
    with open(OUTPUT_TXT, "w", encoding="utf-8") as f:
        # Extract all paragraphs
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text:
                # Check if it's a heading
                if para.style.name.startswith('Heading'):
                    f.write(f"\n{'#' * int(para.style.name[-1]) if para.style.name[-1].isdigit() else '##'} {text}\n\n")
                else:
                    f.write(f"{text}\n")
        
        # Extract tables
        f.write("\n\n=== TABLES ===\n\n")
        for table_idx, table in enumerate(doc.tables):
            f.write(f"\n[TABLE {table_idx + 1}]\n")
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                f.write(" | ".join(row_data) + "\n")
            f.write("\n")

    print(f"Extracted content to {OUTPUT_TXT}")
    print(f"Total paragraphs: {len(doc.paragraphs)}")
    print(f"Total tables: {len(doc.tables)}")

if __name__ == "__main__":
    extract_content()
